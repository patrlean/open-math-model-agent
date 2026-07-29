"""Input ingestion: turn uploaded materials into a workspace the agent can read.

Dispatches by file type and produces:
  - problem.md : extracted prose (PDF/Word text) + a schema summary for every
                 dataset, so the agent understands the task and what data exists
                 without loading whole tables into context.
  - data/*.csv : tabular inputs (Excel sheets, CSVs) normalized to CSV.
  - assets/    : images and figures extracted from documents.

Big tables are NOT dumped into problem.md -- only shape/columns/dtypes/head. The
agent reads full data later, in the sandbox, via code.
"""

from __future__ import annotations

import shutil
from dataclasses import dataclass, field
from pathlib import Path

import pandas as pd

TEXT_EXTS = {".txt", ".md"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
EXCEL_EXTS = {".xlsx", ".xls"}
CSV_EXTS = {".csv"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif"}


@dataclass
class IngestReport:
    problem_md: Path
    data_files: list[str] = field(default_factory=list)
    asset_files: list[str] = field(default_factory=list)
    skipped: list[str] = field(default_factory=list)
    sections: list[str] = field(default_factory=list)  # markdown blocks for problem.md


def _safe_stem(name: str) -> str:
    keep = "".join(c if c.isalnum() or c in "-_" else "_" for c in name)
    return keep.strip("_") or "file"


def _data_summary(df: pd.DataFrame, rel_path: str) -> str:
    cols = ", ".join(f"{c}({df[c].dtype})" for c in df.columns)
    head = df.head(3).to_csv(index=False).strip()
    return (
        f"- file: `{rel_path}`\n"
        f"- shape: {df.shape[0]} rows x {df.shape[1]} cols\n"
        f"- columns: {cols}\n"
        f"- preview (first 3 rows):\n```\n{head}\n```"
    )


def _ingest_pdf(path: Path, workdir: Path, report: IngestReport) -> None:
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    parts, img_count = [], 0
    assets_dir = workdir / "assets"
    for pno, page in enumerate(doc, start=1):
        text = page.get_text().strip()
        if text:
            parts.append(f"[page {pno}]\n{text}")
        for img in page.get_images(full=True):
            xref = img[0]
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.n >= 5:  # CMYK/alpha -> RGB
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                assets_dir.mkdir(exist_ok=True)
                img_count += 1
                rel = f"assets/{_safe_stem(path.stem)}_p{pno}_{img_count}.png"
                pix.save(workdir / rel)
                report.asset_files.append(rel)
            except Exception:
                pass
    doc.close()
    body = "\n\n".join(parts) if parts else "(no extractable text)"
    note = f"\n\n_({img_count} embedded image(s) saved to assets/)_" if img_count else ""
    report.sections.append(f"## From `{path.name}` (PDF)\n\n{body}{note}")


def _ingest_docx(path: Path, workdir: Path, report: IngestReport) -> None:
    import docx  # python-docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for ti, table in enumerate(d.tables, start=1):
        rows = [[c.text for c in row.cells] for row in table.rows]
        if rows:
            df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
            rel = f"data/{_safe_stem(path.stem)}_table{ti}.csv"
            (workdir / "data").mkdir(exist_ok=True)
            df.to_csv(workdir / rel, index=False)
            report.data_files.append(rel)
            parts.append(f"[table {ti} -> {rel}]")
    body = "\n".join(parts) if parts else "(empty document)"
    report.sections.append(f"## From `{path.name}` (Word)\n\n{body}")


def _ingest_excel(path: Path, workdir: Path, report: IngestReport) -> None:
    (workdir / "data").mkdir(exist_ok=True)
    sheets = pd.read_excel(path, sheet_name=None)  # dict of all sheets
    summaries = []
    for sheet, df in sheets.items():
        rel = f"data/{_safe_stem(path.stem)}_{_safe_stem(sheet)}.csv"
        df.to_csv(workdir / rel, index=False)
        report.data_files.append(rel)
        summaries.append(f"### sheet `{sheet}`\n{_data_summary(df, rel)}")
    report.sections.append(
        f"## Data from `{path.name}` (Excel)\n\n" + "\n\n".join(summaries)
    )


def _ingest_csv(path: Path, workdir: Path, report: IngestReport) -> None:
    (workdir / "data").mkdir(exist_ok=True)
    df = pd.read_csv(path)
    rel = f"data/{_safe_stem(path.stem)}.csv"
    df.to_csv(workdir / rel, index=False)
    report.data_files.append(rel)
    report.sections.append(
        f"## Data from `{path.name}` (CSV)\n\n{_data_summary(df, rel)}"
    )


def _ingest_image(path: Path, workdir: Path, report: IngestReport) -> None:
    assets_dir = workdir / "assets"
    assets_dir.mkdir(exist_ok=True)
    rel = f"assets/{_safe_stem(path.stem)}{path.suffix.lower()}"
    shutil.copy2(path, workdir / rel)
    report.asset_files.append(rel)
    # OCR is a later enhancement (needs a system tesseract binary); for now the
    # image is preserved and noted so a vision-capable step can pick it up.
    report.sections.append(
        f"## Image `{path.name}`\n\nSaved to `{rel}` (not yet OCR'd)."
    )


def _ingest_text(path: Path, workdir: Path, report: IngestReport) -> None:
    report.sections.append(
        f"## From `{path.name}`\n\n{path.read_text(errors='replace').strip()}"
    )


_DISPATCH = [
    (PDF_EXTS, _ingest_pdf),
    (DOCX_EXTS, _ingest_docx),
    (EXCEL_EXTS, _ingest_excel),
    (CSV_EXTS, _ingest_csv),
    (IMAGE_EXTS, _ingest_image),
    (TEXT_EXTS, _ingest_text),
]


def ingest(
    paths: list[str | Path],
    workdir: str | Path,
    problem_text: str | None = None,
) -> IngestReport:
    """Ingest a modeling request into ``problem.md`` + ``data/`` + ``assets/``.

    ``problem_text`` is the problem statement entered directly in the chat
    composer. Keeping it in the same document as extracted upload content gives
    the modeling agent one canonical task source regardless of input method.
    """
    workdir = Path(workdir)
    workdir.mkdir(parents=True, exist_ok=True)
    report = IngestReport(problem_md=workdir / "problem.md")

    if problem_text and problem_text.strip():
        report.sections.append(
            "## Problem entered in chat\n\n" + problem_text.strip()
        )

    for raw in paths:
        p = Path(raw)
        ext = p.suffix.lower()
        handler = next((h for exts, h in _DISPATCH if ext in exts), None)
        if handler is None:
            report.skipped.append(str(p))
            continue
        try:
            handler(p, workdir, report)
        except Exception as e:  # one bad file must not abort ingestion
            report.sections.append(f"## From `{p.name}`\n\n[ingest error: {type(e).__name__}: {e}]")

    new_body = "\n\n".join(report.sections) + "\n"
    if report.problem_md.is_file():
        # A follow-up ingest (files added to an already-running conversation):
        # append, never overwrite -- the original problem statement must survive.
        existing = report.problem_md.read_text(errors="replace")
        report.problem_md.write_text(
            existing.rstrip() + "\n\n---\n\n## Additional materials\n\n" + new_body
        )
    else:
        report.problem_md.write_text("# Problem Materials\n\n" + new_body)
    return report
