"""Self-test for the ingestion layer.

Generates sample PDF/Word/Excel/CSV files, ingests them, and prints the resulting
problem.md and data/ listing. Run:  ./.venv/bin/python -m scripts.check_ingest
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import pandas as pd

from mathmodel.ingest.ingest import ingest


def make_samples(d: Path) -> list[Path]:
    paths = []

    # PDF via PyMuPDF
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Problem A: Optimal Sensor Placement\n"
                                "Minimize total coverage cost subject to full coverage.",
                     fontsize=12)
    pdf = d / "problem.pdf"
    doc.save(pdf); doc.close()
    paths.append(pdf)

    # Word via python-docx
    import docx
    wd = docx.Document()
    wd.add_heading("Assumptions", level=1)
    wd.add_paragraph("1. Sensors are independent. 2. Cost is additive.")
    t = wd.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "sensor"; t.rows[0].cells[1].text = "cost"
    t.rows[1].cells[0].text = "S1"; t.rows[1].cells[1].text = "10"
    docx_path = d / "notes.docx"
    wd.save(docx_path)
    paths.append(docx_path)

    # Excel (2 sheets) via pandas
    xlsx = d / "measurements.xlsx"
    with pd.ExcelWriter(xlsx) as xl:
        pd.DataFrame({"x": [1, 2, 3], "y": [2.1, 4.0, 6.2]}).to_excel(xl, sheet_name="run1", index=False)
        pd.DataFrame({"t": [0, 1], "temp": [20, 22]}).to_excel(xl, sheet_name="run2", index=False)
    paths.append(xlsx)

    # CSV
    csv = d / "extra.csv"
    pd.DataFrame({"id": [1, 2], "val": [0.5, 0.7]}).to_csv(csv, index=False)
    paths.append(csv)

    return paths


def main() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = tmp / "uploads"; src.mkdir()
        samples = make_samples(src)
        workdir = tmp / "run"

        report = ingest(samples, workdir)

        print("=== IngestReport ===")
        print("data_files:", report.data_files)
        print("asset_files:", report.asset_files)
        print("skipped:", report.skipped)
        print("\n=== problem.md ===")
        print(report.problem_md.read_text())
        print("=== data/ contents ===")
        for f in report.data_files:
            print(f"\n--- {f} ---")
            print((workdir / f).read_text().strip())

        assert len(report.data_files) >= 3, "expected excel(2)+csv+docx table"
        assert "Optimal Sensor Placement" in report.problem_md.read_text()
        print("\nOK: ingestion produced problem.md + data files.")


if __name__ == "__main__":
    main()
